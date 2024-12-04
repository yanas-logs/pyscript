from docx import Document

# Create a new document
doc = Document()

# Title
doc.add_heading('Lorem Ipsum Dolor Sit Amet', level=1)

# Introduction
doc.add_paragraph(
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit. Praesent vehicula dapibus efficitur. "
    "Nam at leo eget augue blandit faucibus quis nec erat. Phasellus in est magna."
)

# Table title
doc.add_heading('Analysis Table Dummy', level=2)

# Create a table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Add headers
header_cells = table.rows[0].cells
header_cells[0].text = "No."
header_cells[1].text = "Parameter/Item"
header_cells[2].text = "Description"
header_cells[3].text = "Status"
header_cells[4].text = "Comment"

# Data Table
data = [
    ["1", "Lorem Ipsum", "Dolor Sit Amet", "Active", "Lorem ipsum dolor sit amet, consectetur adipiscing elit."],
    ["2", "Sed Do Eiusmod", "Tempor Incididunt", "Non-active", "Ut labore et dolore magna aliqua."],
    ["3", "Ut Enim Ad Minim", "Veniam", "Process", "Quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat."],
    ["4", "Duis Aute Irure", "Dolor in Reprehenderit", "Pending", "In voluptate velit esse cillum dolore eu fugiat nulla pariatur.\n\n"]
]

# Add data to the table
for row_data in data:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = cell_data

# Analysis
doc.add_heading('Analisis Dummy', level=2)
doc.add_paragraph(
    "1. Lorem ipsum dolor sit amet, consectetur adipiscing elit. Nulla cursus nulla in quam cursus, "
    "vitae pellentesque tortor sagittis.\n\n"
    "2. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, "
    "quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.\n\n"
    "3. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.\n\n"
    "4. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum."
)

# Lessons
doc.add_heading('Lessons from Dummy', level=2)
doc.add_paragraph(
    "1. Proin volutpat, felis non convallis vehicula, lacus tortor vestibulum lacus, "
    "in viverra erat erat et eros.\n"
    "2. Vestibulum ante ipsum primis in faucibus orci luctus et ultrices posuere cubilia curae.\n"
    "3. Suspendisse potenti. Maecenas vehicula massa nec sapien vehicula, non interdum ligula dapibus.\n"
    "4. Aliquam erat volutpat. Aenean id tortor vel eros consectetur fermentum in eu arcu."
)

# Save the document to a file
file_path = "/mnt/data/result.docx"  # Ensure the path is valid for your system
doc.save(file_path)

file_path
